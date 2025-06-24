import os
import json
import requests
from typing import List, Dict, Any, Optional
from simple_vectordb import SimpleVectorDB
from model_config import ModelConfig
import hashlib
import fitz  # PyMuPDF
import docx
import tiktoken


class OllamaEmbeddings:
    def __init__(self, model: str = "nomic-embed-text", base_url: str = "http://localhost:11434"):
        self.model = model
        self.base_url = base_url
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        embeddings = []
        for text in texts:
            response = requests.post(
                f"{self.base_url}/api/embeddings",
                json={"model": self.model, "prompt": text}
            )
            if response.status_code == 200:
                embeddings.append(response.json()["embedding"])
            else:
                raise Exception(f"Failed to get embedding: {response.text}")
        return embeddings
    
    def embed_query(self, text: str) -> List[float]:
        return self.embed_documents([text])[0]


class DocumentProcessor:
    def __init__(self, chunk_size: int = 3000, chunk_overlap: int = 400):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    def load_document(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext == '.docx':
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _extract_pdf(self, file_path: str) -> str:
        text = ""
        try:
            doc = fitz.open(file_path)
            for page_num in range(doc.page_count):
                page = doc[page_num]
                # Try text extraction first (fast)
                page_text = page.get_text()
                if not page_text.strip():
                    # Try OCR if available
                    try:
                        pix = page.get_pixmap()
                        ocr_text = page.get_textpage_ocr().extractText()
                        page_text = ocr_text
                    except Exception as ocr_e:
                        print(f"Warning: OCR failed for page {page_num + 1} in {file_path}: {ocr_e}")
                        print("Note: Install tesseract-ocr for image-based PDF support")
                        page_text = ""
                text += page_text + "\n"
            doc.close()
        except Exception as e:
            print(f"Warning: Could not extract text from {file_path}: {e}")
        return text
    
    def _extract_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        # Extract headers and footers from each section
        for section in doc.sections:
            # Extract headers
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(f"[HEADER] {paragraph.text}")
            
            # Extract footers
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(f"[FOOTER] {paragraph.text}")
        
        extracted_text = "\n".join(full_text)
        print(f"Extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
    
    def chunk_text(self, text: str) -> List[str]:
        tokens = self.encoding.encode(text)
        chunks = []
        
        for i in range(0, len(tokens), self.chunk_size - self.chunk_overlap):
            chunk_tokens = tokens[i:i + self.chunk_size]
            chunk_text = self.encoding.decode(chunk_tokens)
            chunks.append(chunk_text)
        
        return chunks


class RAGSystem:
    def __init__(self, persist_directory: str = "./vector_db", llm_model: str = "auto", conversation_memory_size: int = 5):
        self.persist_directory = persist_directory
        self.vectordb = SimpleVectorDB(persist_directory)
        self.embeddings = OllamaEmbeddings()
        self.processor = DocumentProcessor()
        self.llm_url = "http://localhost:11434/api/generate"
        
        # Simple conversation memory
        self.conversation_history = []
        self.conversation_memory_size = conversation_memory_size
        
        # Auto-select model based on system resources
        if llm_model == "auto":
            self.llm_model = ModelConfig.get_recommended_model('balanced')
            print(f"Auto-selected model: {self.llm_model}")
        else:
            self.llm_model = llm_model
    
    def add_document(self, file_path: str) -> None:
        print(f"Processing document: {file_path}")
        
        # Load and process document
        text = self.processor.load_document(file_path)
        chunks = self.processor.chunk_text(text)
        
        # Generate embeddings
        embeddings = self.embeddings.embed_documents(chunks)
        
        # Create unique IDs for chunks
        doc_id = hashlib.md5(file_path.encode()).hexdigest()
        ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
        
        # Prepare metadata
        metadatas = [{"source": file_path, "chunk_id": i} for i in range(len(chunks))]
        
        # Add to vector database
        self.vectordb.add(
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadatas,
            ids=ids
        )
        
        print(f"Added {len(chunks)} chunks from {file_path}")
    
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """Analyze document extraction and chunking without adding to vector DB"""
        print(f"Analyzing document: {file_path}")
        
        # Load and process document
        text = self.processor.load_document(file_path)
        chunks = self.processor.chunk_text(text)
        
        # Calculate statistics
        total_chars = len(text)
        total_tokens = len(self.processor.encoding.encode(text))
        chunk_sizes = [len(chunk) for chunk in chunks]
        chunk_tokens = [len(self.processor.encoding.encode(chunk)) for chunk in chunks]
        
        analysis = {
            'file_path': file_path,
            'total_characters': total_chars,
            'total_tokens': total_tokens,
            'num_chunks': len(chunks),
            'avg_chunk_chars': sum(chunk_sizes) / len(chunk_sizes) if chunks else 0,
            'avg_chunk_tokens': sum(chunk_tokens) / len(chunk_tokens) if chunks else 0,
            'min_chunk_tokens': min(chunk_tokens) if chunks else 0,
            'max_chunk_tokens': max(chunk_tokens) if chunks else 0,
            'chunk_token_distribution': chunk_tokens
        }
        
        print(f"Document Analysis:")
        print(f"  File: {file_path}")
        print(f"  Total characters: {total_chars:,}")
        print(f"  Total tokens: {total_tokens:,}")
        print(f"  Number of chunks: {len(chunks)}")
        print(f"  Average chunk size: {analysis['avg_chunk_tokens']:.0f} tokens")
        print(f"  Chunk size range: {analysis['min_chunk_tokens']}-{analysis['max_chunk_tokens']} tokens")
        
        return analysis
    
    def search(self, query: str, n_results: int = 15) -> List[Dict[str, Any]]:
        # Generate query embedding
        query_embedding = self.embeddings.embed_query(query)
        
        # Search in vector database
        results = self.vectordb.query(query_embedding, n_results)
        
        # Format results
        search_results = []
        for i in range(len(results['documents'][0])):
            search_results.append({
                'content': results['documents'][0][i],
                'metadata': results['metadatas'][0][i],
                'distance': results['distances'][0][i]
            })
        
        return search_results
    
    def generate_response(self, query: str, context: str, model: str = None, timeout: int = 600, use_conversation: bool = False) -> str:
        if model is None:
            model = self.llm_model
        
        # Build conversation context if available
        conversation_context = ""
        if use_conversation and self.conversation_history:
            conversation_context = "\nPrevious conversation:\n"
            for i, exchange in enumerate(self.conversation_history):
                conversation_context += f"Q{i+1}: {exchange['question']}\nA{i+1}: {exchange['answer']}\n"
            conversation_context += "\n"
        
        prompt = f"""Answer this question using the context provided. If there's previous conversation, consider it for continuity.{conversation_context}
Document context: {context}

Current question: {query}

Answer:"""
        
        try:
            response = requests.post(
                self.llm_url,
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,
                        "top_p": 0.9,
                        "num_predict": 1500
                    }
                },
                timeout=timeout
            )
            
            if response.status_code == 200:
                return response.json()["response"]
            else:
                raise Exception(f"Failed to generate response: {response.text}")
        except requests.exceptions.Timeout:
            return "Response timed out. The model may be overloaded or the query too complex."
        except Exception as e:
            return f"Error generating response: {str(e)}"
    
    def query(self, question: str, n_results: int = 15, debug: bool = False, timeout: int = 600, use_conversation: bool = False) -> Dict[str, Any]:
        # Search for relevant documents
        search_results = self.search(question, n_results)
        
        # Prepare context from search results
        context = "\n\n".join([result['content'] for result in search_results])
        
        # Generate response using LLM with conversation context if enabled
        response = self.generate_response(question, context, timeout=timeout, use_conversation=use_conversation)
        
        # Update conversation history if using conversation mode
        if use_conversation:
            self._update_conversation_history(question, response)
        
        result = {
            'question': question,
            'answer': response,
            'sources': [result['metadata']['source'] for result in search_results],
            'context_chunks': len(search_results),
            'conversation_mode': use_conversation
        }
        
        if debug:
            print(f"\n=== DEBUG: Retrieved {len(search_results)} chunks ===")
            for i, chunk_result in enumerate(search_results):
                print(f"\nChunk {i+1} (distance: {chunk_result['distance']:.4f}):")
                print(f"Source: {chunk_result['metadata']['source']}")
                print(f"Content preview: {chunk_result['content'][:200]}...")
                print(f"Full content length: {len(chunk_result['content'])} characters")
            
            print(f"\n=== Total context length: {len(context)} characters ===")
            
            if use_conversation and self.conversation_history:
                print(f"\n=== Conversation History ({len(self.conversation_history)} exchanges) ===")
                for i, exchange in enumerate(self.conversation_history):
                    print(f"Exchange {i+1}: {exchange['question'][:100]}...")
            
            result['debug_info'] = {
                'retrieved_chunks': search_results,
                'total_context_length': len(context),
                'conversation_length': len(self.conversation_history)
            }
        
        return result
    
    def _update_conversation_history(self, question: str, answer: str) -> None:
        """Update conversation history with FIFO buffer"""
        self.conversation_history.append({
            'question': question,
            'answer': answer
        })
        
        # Keep only the last N exchanges
        if len(self.conversation_history) > self.conversation_memory_size:
            self.conversation_history.pop(0)
    
    def reset_conversation(self) -> None:
        """Clear conversation history"""
        self.conversation_history = []
        print("Conversation history cleared.")
    
    def show_conversation(self) -> None:
        """Display current conversation history"""
        if not self.conversation_history:
            print("No conversation history.")
            return
        
        print(f"\n=== Conversation History ({len(self.conversation_history)} exchanges) ===")
        for i, exchange in enumerate(self.conversation_history):
            print(f"\nQ{i+1}: {exchange['question']}")
            print(f"A{i+1}: {exchange['answer'][:200]}..." if len(exchange['answer']) > 200 else f"A{i+1}: {exchange['answer']}")
        print(f"\nMemory capacity: {len(self.conversation_history)}/{self.conversation_memory_size}")
    
    def get_conversation_info(self) -> Dict[str, Any]:
        """Get conversation statistics"""
        return {
            'history_length': len(self.conversation_history),
            'memory_capacity': self.conversation_memory_size,
            'has_conversation': len(self.conversation_history) > 0
        }
    
    def remove_document(self, file_path: str) -> int:
        """Remove a document and all its chunks from the knowledge base"""
        removed_count = self.vectordb.remove_by_source(file_path)
        if removed_count > 0:
            print(f"Removed {removed_count} chunks from {file_path}")
        else:
            print(f"No chunks found for {file_path}")
        return removed_count
    
    def list_documents(self) -> List[str]:
        """List all documents in the knowledge base"""
        return self.vectordb.list_sources()
    
    def get_document_info(self) -> Dict[str, int]:
        """Get information about documents and their chunk counts"""
        return self.vectordb.get_source_info()
    
    def show_documents(self) -> None:
        """Display all documents in the knowledge base with their chunk counts"""
        doc_info = self.get_document_info()
        if not doc_info:
            print("No documents in knowledge base.")
            return
        
        print(f"\n=== Knowledge Base Documents ({len(doc_info)} files) ===")
        total_chunks = 0
        for source, chunk_count in doc_info.items():
            print(f"{chunk_count:3d} chunks: {source}")
            total_chunks += chunk_count
        print(f"\nTotal: {total_chunks} chunks across {len(doc_info)} documents")
    
    def clear_knowledge_base(self) -> int:
        """Remove all documents from the knowledge base"""
        removed_count = self.vectordb.clear_all()
        if removed_count > 0:
            print(f"Cleared knowledge base: removed {removed_count} chunks")
        else:
            print("Knowledge base was already empty")
        return removed_count